"""
Stage 1: File Parser
Converts PDF or DOCX manuscripts into a list of Chapter objects.
Detects chapter boundaries from heading styles (DOCX) or bookmarks/regex (PDF).
"""

import re
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document

from framework.models.chunk import Chapter

# Patterns for chapter detection fallback (when no headings/bookmarks found)
# Strict patterns are tried first; the loose numbered pattern is only used
# if strict patterns find too few chapters.
STRICT_CHAPTER_PATTERNS = [
    re.compile(r"^(Chapter|CHAPTER)\s+[IVXLCDM\d]+", re.IGNORECASE),
    re.compile(r"^(Part|PART)\s+[IVXLCDM\d]+", re.IGNORECASE),
    re.compile(r"^(Book|BOOK)\s+[IVXLCDM\d]+", re.IGNORECASE),
    re.compile(r"^(Section|SECTION)\s+\d+", re.IGNORECASE),
]
LOOSE_CHAPTER_PATTERNS = [
    re.compile(r"^\d+\.\s+[A-Z]"),  # "1. Title" format — aggressive, used as last resort
]
CHAPTER_PATTERNS = STRICT_CHAPTER_PATTERNS + LOOSE_CHAPTER_PATTERNS

# Maximum expected chapters before falling back to strict-only patterns
MAX_REASONABLE_CHAPTERS = 80

# DOCX heading styles that indicate chapter-level breaks
CHAPTER_HEADING_STYLES = {"Heading 1", "Heading 2", "heading 1", "heading 2"}

# Minimum word count for a chapter to be valid (avoids capturing blank headings)
MIN_CHAPTER_WORDS = 50


def parse_file(file_path: str | Path) -> list[Chapter]:
    """Parse a manuscript file into chapters. Dispatches by extension."""
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".docx":
        return _parse_docx(path)
    elif ext == ".pdf":
        return _parse_pdf(path)
    elif ext == ".txt":
        return _parse_txt(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Use PDF, DOCX, or TXT.")


def _parse_docx(path: Path) -> list[Chapter]:
    """Parse DOCX using heading styles for chapter detection."""
    doc = Document(str(path))

    chapters: list[Chapter] = []
    current_title = ""
    current_text_parts: list[str] = []
    chapter_num = 0

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()

        if not text:
            if current_text_parts:
                current_text_parts.append("")  # preserve paragraph breaks
            continue

        is_heading = style_name in CHAPTER_HEADING_STYLES

        if is_heading:
            # Save previous chapter if it has content
            if current_text_parts and _word_count(current_text_parts) >= MIN_CHAPTER_WORDS:
                chapter_num += 1
                chapters.append(Chapter(
                    number=chapter_num,
                    title=current_title or f"Section {chapter_num}",
                    raw_text="\n".join(current_text_parts).strip(),
                ))
            current_title = text
            current_text_parts = []
        else:
            current_text_parts.append(text)

    # Capture final chapter
    if current_text_parts and _word_count(current_text_parts) >= MIN_CHAPTER_WORDS:
        chapter_num += 1
        chapters.append(Chapter(
            number=chapter_num,
            title=current_title or f"Section {chapter_num}",
            raw_text="\n".join(current_text_parts).strip(),
        ))

    # If heading-based detection found too few chapters, try regex fallback
    if len(chapters) <= 1:
        fallback = _parse_docx_regex_fallback(doc)
        if len(fallback) > len(chapters):
            return fallback

    return chapters


def _parse_docx_regex_fallback(doc: Document) -> list[Chapter]:
    """Fallback: detect chapters via regex patterns in paragraph text."""
    chapters: list[Chapter] = []
    current_title = ""
    current_text_parts: list[str] = []
    chapter_num = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            if current_text_parts:
                current_text_parts.append("")
            continue

        if _matches_chapter_pattern(text):
            if current_text_parts and _word_count(current_text_parts) >= MIN_CHAPTER_WORDS:
                chapter_num += 1
                chapters.append(Chapter(
                    number=chapter_num,
                    title=current_title or f"Section {chapter_num}",
                    raw_text="\n".join(current_text_parts).strip(),
                ))
            current_title = text
            current_text_parts = []
        else:
            current_text_parts.append(text)

    if current_text_parts and _word_count(current_text_parts) >= MIN_CHAPTER_WORDS:
        chapter_num += 1
        chapters.append(Chapter(
            number=chapter_num,
            title=current_title or f"Section {chapter_num}",
            raw_text="\n".join(current_text_parts).strip(),
        ))

    return chapters


def _parse_pdf(path: Path) -> list[Chapter]:
    """Parse PDF using bookmarks for chapter detection, falling back to regex."""
    doc = fitz.open(str(path))

    # Try bookmark-based detection first
    toc = doc.get_toc()  # list of [level, title, page_number]
    if toc:
        return _parse_pdf_from_toc(doc, toc)

    # Fallback: extract all text and split by regex
    return _parse_pdf_regex_fallback(doc)


def _parse_pdf_from_toc(doc: fitz.Document, toc: list) -> list[Chapter]:
    """Parse PDF chapters using table of contents / bookmarks."""
    # Filter to top-level entries (level 1 or 2)
    top_entries = [e for e in toc if e[0] <= 2]
    if not top_entries:
        top_entries = toc

    chapters: list[Chapter] = []
    for i, entry in enumerate(top_entries):
        _level, title, start_page = entry
        start_page -= 1  # 0-indexed

        # End page is start of next entry, or end of document
        if i + 1 < len(top_entries):
            end_page = top_entries[i + 1][2] - 1
        else:
            end_page = len(doc)

        text_parts = []
        for page_num in range(start_page, min(end_page, len(doc))):
            page_text = doc[page_num].get_text()
            if page_text.strip():
                text_parts.append(page_text.strip())

        raw_text = "\n".join(text_parts)
        if len(raw_text.split()) >= MIN_CHAPTER_WORDS:
            chapters.append(Chapter(
                number=len(chapters) + 1,
                title=title.strip(),
                raw_text=raw_text,
            ))

    return chapters


def _parse_pdf_regex_fallback(doc: fitz.Document) -> list[Chapter]:
    """Fallback: extract all PDF text and split by chapter regex patterns."""
    full_text = ""
    for page in doc:
        full_text += page.get_text() + "\n"

    return _split_text_by_regex(full_text)


def _parse_txt(path: Path) -> list[Chapter]:
    """Parse plain text file, splitting by regex patterns."""
    text = path.read_text(encoding="utf-8")
    return _split_text_by_regex(text)


def _split_text_by_regex(text: str) -> list[Chapter]:
    """Split raw text into chapters using regex chapter patterns.
    Tries strict patterns first; only adds loose patterns if strict yields too few."""
    # Try strict patterns first
    chapters = _split_with_patterns(text, STRICT_CHAPTER_PATTERNS)

    # If strict found too few, try all patterns
    if len(chapters) <= 1:
        all_chapters = _split_with_patterns(text, CHAPTER_PATTERNS)
        # Only use loose results if they produce a reasonable count
        if 1 < len(all_chapters) <= MAX_REASONABLE_CHAPTERS:
            chapters = all_chapters

    # If still nothing, treat whole document as one chapter
    if not chapters and len(text.split()) >= MIN_CHAPTER_WORDS:
        chapters = [Chapter(number=1, title="Full Document", raw_text=text.strip())]

    return chapters


def _split_with_patterns(text: str, patterns: list) -> list[Chapter]:
    """Split text using a specific set of regex patterns."""
    lines = text.split("\n")
    chapters: list[Chapter] = []
    current_title = ""
    current_lines: list[str] = []
    chapter_num = 0

    for line in lines:
        stripped = line.strip()
        if _matches_patterns(stripped, patterns):
            if current_lines and _word_count(current_lines) >= MIN_CHAPTER_WORDS:
                chapter_num += 1
                chapters.append(Chapter(
                    number=chapter_num,
                    title=current_title or f"Section {chapter_num}",
                    raw_text="\n".join(current_lines).strip(),
                ))
            current_title = stripped
            current_lines = []
        else:
            current_lines.append(line)

    if current_lines and _word_count(current_lines) >= MIN_CHAPTER_WORDS:
        chapter_num += 1
        chapters.append(Chapter(
            number=chapter_num,
            title=current_title or f"Section {chapter_num}",
            raw_text="\n".join(current_lines).strip(),
        ))

    return chapters


def _matches_chapter_pattern(text: str) -> bool:
    """Check if a line matches any chapter heading pattern."""
    return _matches_patterns(text, CHAPTER_PATTERNS)


def _matches_patterns(text: str, patterns: list) -> bool:
    """Check if a line matches any of the given regex patterns."""
    if len(text) > 200:  # headings are short
        return False
    return any(p.match(text) for p in patterns)


def _word_count(parts: list[str]) -> int:
    """Count total words in a list of text parts."""
    return sum(len(p.split()) for p in parts)
